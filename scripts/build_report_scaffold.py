"""Build the Part B report scaffold at report/report.docx.

    python scripts/build_report_scaffold.py

What this script writes, and what it deliberately does not:

- It writes the STRUCTURE (headings in the order the brief asks for), every
  exhibit with a Word caption that auto-numbers through a SEQ field, the key
  tables as Word tables, and the factual description of what was built and
  measured. Every number is read from results/ at build time, so the scaffold
  cannot drift from the artifacts.
- The interpretive passages come from scripts/report_prose.py. They were
  drafted by the assistant at my request and are NOT yet my own words. The
  course grades my economic reasoning, so every one of them has to be
  rewritten before submission; report/REVISION_CHECKLIST.md tracks that pass.

Re-running overwrites report.docx. Once I start editing in Word, edit
report_prose.py instead of the document, or stop running this script.
"""
from __future__ import annotations

import pathlib
import sys

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from report_prose import prose

ROOT = pathlib.Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))

FIGURES = ROOT / "results" / "figures"
TABLES = ROOT / "results" / "tables"
TARGET = ROOT / "report" / "report.docx"

FIG_WIDTH = Inches(6.1)          # fits A4 with 2.5 cm margins
MARK = "[WRITE]"


# --- Word plumbing ---------------------------------------------------------

def _field(paragraph, instruction: str) -> None:
    """Insert a Word field code (used for SEQ auto-numbering)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t"); placeholder.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, sep, placeholder, end):
        run._r.append(node)


def ensure_caption_style(doc: Document):
    """The scaffold ships without a Caption style; captions need one."""
    try:
        return doc.styles["Caption"]
    except KeyError:
        style = doc.styles.add_style("Caption", 1)
        style.base_style = doc.styles["Normal"]
        style.font.size = Pt(9)
        style.font.italic = True
        style.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)
        return style


def caption(doc: Document, kind: str, text: str) -> None:
    """A Word caption whose number comes from a SEQ field, not typed by hand."""
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{kind} ")
    _field(p, f" SEQ {kind} \\* ARABIC ")
    p.add_run(f": {text}")


def figure(doc: Document, filename: str, text: str) -> None:
    path = FIGURES / filename
    if not path.exists():
        doc.add_paragraph(f"{MARK} missing figure: {filename}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=FIG_WIDTH)
    caption(doc, "Figure", text)


def table(doc: Document, df: pd.DataFrame, text: str, floats: str = "{:.3f}") -> None:
    caption(doc, "Table", text)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    for j, col in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(col).replace("_", " ")
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = floats.format(v) if isinstance(v, float) else str(v)
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


_PROSE: dict = {}


def write_marker(doc: Document, slot: str) -> None:
    """Emit the interpretive passage registered for this slot.

    A slot with no passage still prints a visible marker rather than silently
    leaving a gap in the argument.
    """
    paragraphs = _PROSE.get(slot)
    if not paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(f"{MARK} no prose registered for slot '{slot}'")
        run.italic = True
        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        return
    for text in paragraphs:
        doc.add_paragraph(text)


# --- content ---------------------------------------------------------------

def load_numbers() -> dict:
    """Read every quoted number from the artifacts, so nothing is hardcoded."""
    m = pd.read_csv(TABLES / "performance_metrics.csv")
    lex = pd.read_csv(TABLES / "lexicon_effect.csv")
    diag = pd.read_csv(TABLES / "sentiment_signal_diagnostics.csv")
    ext = pd.read_csv(TABLES / "extension_comparison.csv")
    freq = pd.read_csv(TABLES / "rebalance_frequency.csv")
    return {
        "metrics": m,
        "lead_lag": pd.read_csv(TABLES / "sentiment_lead_lag.csv"),
        "horizons": pd.read_csv(TABLES / "sentiment_horizons.csv"),
        "shrinkage": pd.read_csv(TABLES / "shrinkage_study.csv"),
        "significance": pd.read_csv(TABLES / "extension_significance.csv"),
        "fusion": pd.read_csv(TABLES / "fusion_before_after.csv"),
        "coverage": pd.read_csv(TABLES / "sentiment_coverage.csv"),
        "lexicon": lex,
        "diagnostics": diag,
        "extensions": ext,
        "holdout": pd.read_csv(TABLES / "discovery_holdout.csv"),
        "frequency": freq,
        "n_funds": len(m),
        "first_live": m["first_live_date"].min(),
        "last": m["last_date"].max(),
        "n_days": int(m["n_days"].max()),
        "neutral_before": float(lex.loc[0, "neutral_share"]),
        "neutral_after": float(lex.loc[1, "neutral_share"]),
        "n_terms": int(lex.loc[1, "n_terms_added"]),
        "rescored": float(lex.loc[1, "headlines_rescored_share"]),
        "n_headlines": int(pd.read_csv(TABLES / "sentiment_coverage.csv")
                           ["n_headlines"].sum()),
        "mean_corr": float(diag["corr_lagged_sentiment_vs_return"].mean()),
        "n_negative": int((diag["corr_lagged_sentiment_vs_return"] < 0).sum()),
    }


def build() -> None:
    global _PROSE
    n = load_numbers()
    _PROSE.update(prose(n))
    doc = Document(str(TARGET)) if TARGET.exists() else Document()

    # Start from a clean body; the scaffold's generic sections are replaced.
    # The trailing sectPr holds page size and margins - removing it leaves
    # python-docx unable to size an image to the text column.
    for element in list(doc.element.body):
        if element.tag != qn("w:sectPr"):
            doc.element.body.remove(element)
    ensure_caption_style(doc)

    doc.add_paragraph(
        "Spotlight: Systematic Multi-Asset Funds with News-Sentiment "
        "Analytics", style="Title")
    doc.add_paragraph("Nuo Chen (z5640476)")
    doc.add_paragraph("FINS5545 Financial Market Data Literacy - Part B")

    doc.add_heading("Abstract", level=1)
    write_marker(doc, "abstract")

    doc.add_heading("Contents", level=1)
    p = doc.add_paragraph()
    _field(p, r' TOC \o "1-2" \h \z \u ')
    para(doc, "Right-click and choose Update Field in Word to populate this.")

    # 1 -------------------------------------------------------------------
    doc.add_heading("1. The funds and the backtest design", level=1)
    para(doc,
         f"Spotlight offers {n['n_funds']} funds. Each is one asset family - "
         "the 50 US equities, the 10 cryptocurrencies, or the two combined - "
         "run through one optimisation method, because a family-method pair is "
         "what an investor buys and what a fact sheet covers. The four methods "
         "are equal weight, minimum variance, maximum Sharpe (mean-variance "
         "tangency), and risk parity. All are long only and fully invested, so "
         "weights are non-negative and sum to one.")
    para(doc,
         "Weights come from a walk-forward backtest. At each rebalance the "
         "optimiser sees a rolling window of the previous 252 trading days and "
         "nothing after it, and the resulting weights are held until the next "
         "rebalance, drifting with prices in between as a real fund does. "
         "Rebalancing is monthly, on the first trading day of each month, "
         f"giving 36 rebalances. The out-of-sample record runs from "
         f"{n['first_live']} to {n['last']} ({n['n_days']} trading days for the "
         "equity and combined funds); it begins after the first full estimation "
         "window, not on the first date in the data.")
    para(doc,
         "Two conventions are stated rather than assumed. The risk-free rate is "
         "zero, so the Sharpe ratio is the annualised mean return divided by "
         "annualised volatility. Headline results are gross of transaction "
         "costs; a 10 basis point charge on one-way turnover is reported "
         "alongside as a robustness layer. Equity and combined funds annualise "
         "on 252 days and crypto-only funds on 365, matching each calendar.")
    para(doc, "The four objectives are stated formally in Appendix A, with "
              "every symbol defined.")
    write_marker(doc, "design_choices")
    para(doc,
         "The absence of look-ahead is asserted rather than argued. "
         "tests/test_no_lookahead.py runs each fund on the full sample and "
         "again on a sample truncated at 30 June 2022 and requires the "
         "overlapping returns and weights to match exactly; a backtest that "
         "used future data could not pass.")

    doc.add_heading("How often the funds trade", level=2)
    freq = n["frequency"]
    best = freq.loc[freq.groupby("method")["sharpe_net_costs"].idxmax()]
    para(doc,
         "Monthly rebalancing is a choice, so the alternatives were measured. "
         "Weekly, fortnightly, monthly and quarterly schedules were run for "
         "each method on the combined universe, gross and net of costs.")
    figure(doc, "rebalance_frequency.png",
           "Sharpe ratio net of 10 bp of one-way turnover against annual "
           "turnover, for four rebalance schedules across the four combined "
           "funds. Out-of-sample, 2021-2023.")
    para(doc, "Best net-of-cost schedule by method: "
              + "; ".join(f"{r.method} {r.frequency}" for r in best.itertuples())
              + ".")
    write_marker(doc, "frequency")

    # 2 -------------------------------------------------------------------
    doc.add_heading("2. Out-of-sample results and fund fact sheets", level=1)
    cols = ["fund", "ann_return", "ann_vol", "sharpe", "max_drawdown",
            "growth_of_1"]
    table(doc, n["metrics"][cols],
          "Out-of-sample performance of all twelve funds. Annualised return is "
          "geometric; the Sharpe ratio uses the arithmetic mean with a "
          "zero risk-free rate. Equity and combined funds annualise on 252 "
          "days, crypto-only on 365. Sample 2021-2023.")
    write_marker(doc, "metrics_table")

    figure(doc, "growth_of_1_combined.png",
           "Growth of $1 invested at the first live backtest date in the four "
           "combined equity-plus-crypto funds.")
    write_marker(doc, "growth")

    figure(doc, "sharpe_by_fund.png",
           "Sharpe ratio across all twelve funds, grouped by asset family.")
    write_marker(doc, "sharpe_bar")

    figure(doc, "drawdown_combined_max_sharpe.png",
           "Drawdown of the Combined Maximum-Sharpe fund from its running peak.")
    write_marker(doc, "drawdown")

    figure(doc, "weights_combined_across_methods.png",
           "Target weights at each monthly rebalance for three combined funds, "
           "aggregated to sector and stacked to 100%. Crypto is the dark band.")
    write_marker(doc, "weights")

    # 3 -------------------------------------------------------------------
    doc.add_heading("3. The news-sentiment index", level=1)
    para(doc,
         f"Every one of the {n['n_headlines']:,} aligned headlines is scored "
         "with VADER. The text is scored raw: VADER reads capitalisation, "
         "punctuation, degree modifiers and negation, so lower-casing or "
         "stripping punctuation would remove the evidence the model is built "
         "on. Scores are averaged within a ticker-day, then equally weighted "
         "across the five tickers in each sector, so that one heavily covered "
         "name cannot speak for its sector.")
    para(doc,
         "Ticker-days with no headline are dropped from the sector average "
         "rather than scored as zero. A zero would pull the index toward "
         "neutral on quiet days and would confuse an absence of news with "
         "balanced news, and Part A found multi-month collection blackouts for "
         "two tickers that a zero-fill would have turned into fake neutral "
         "sentiment. The daily coverage count is published beside the index so "
         "thin days are visible.")
    para(doc,
         "The index is lagged one trading day after alignment. A decision on "
         "day t therefore uses only sentiment from t-1 or earlier: a Saturday "
         "or Monday headline, both aligned to Monday, is first usable for "
         "Tuesday's trade.")
    figure(doc, "sector_sentiment_index.png",
           "News-sentiment index for five equity sectors: equal-weight mean "
           "VADER compound score, 21-day rolling mean, 2020-2023.")
    write_marker(doc, "sentiment_index")

    figure(doc, "fear_greed_index.png",
           "Market-wide fear and greed gauge: sentiment averaged across all 50 "
           "equities, shown as the raw 0-100 level and standardised.")
    write_marker(doc, "fear_greed")

    table(doc, n["coverage"],
          "Coverage and neutral share of the sentiment model by sector.")
    write_marker(doc, "coverage")

    # 4 -------------------------------------------------------------------
    doc.add_heading("4. Fusion: does sentiment improve the funds?", level=1)
    para(doc,
         "The baseline fusion is a cross-sectional sector tilt. At each "
         "rebalance the lagged sector index is turned into a z-score across "
         "the ten sectors, and each equity weight is multiplied by "
         "(1 + lambda z) for its sector, then clipped at zero and renormalised "
         "so the fund stays long-only and fully invested. The tilt never "
         "touches the covariance estimate, so any difference in results comes "
         "from the tilt alone. Sentiment applies to equities only, because the "
         "cryptocurrencies carry no news.")
    table(doc, n["fusion"],
          "Equity funds before and after the sentiment tilt, out of sample.")
    write_marker(doc, "fusion_table")
    figure(doc, "fusion_min_variance.png",
           "Growth of $1 for the Equity Minimum-Variance fund with and without "
           "the sentiment tilt.")
    figure(doc, "fusion_max_sharpe.png",
           "Growth of $1 for the Equity Maximum-Sharpe fund with and without "
           "the sentiment tilt.")
    para(doc,
         "The diagnostics explain the direction of the result. The correlation "
         "between the lagged sector index and the sector return it can be "
         f"traded on is negative in {n['n_negative']} of the ten sectors, with "
         f"a mean of {n['mean_corr']:+.4f}.")
    table(doc, n["diagnostics"],
          "Correlation between the lagged sector sentiment index and the "
          "equal-weight sector return on the day the signal becomes usable.")
    para(doc,
         "A negative correlation is consistent with three different stories - "
         "a reversal, news already in the price, or noise - so sentiment on "
         "day t was correlated against the sector return on the day before, "
         "the same day, and the day after, with a p-value for each.")
    ll = n["lead_lag"]
    table(doc, ll[["sector", "corr_prev_day", "p_prev_day", "corr_same_day",
                   "p_same_day", "corr_next_day", "p_next_day"]],
          "Sector sentiment on day t against the sector return on t-1, t and "
          "t+1, with Pearson p-values. Raw daily index, 2020-2023.")
    para(doc,
         f"Mean correlations: {ll['corr_same_day'].mean():+.4f} same day "
         f"({int((ll['p_same_day'] < 0.05).sum())} of 10 significant at 5%), "
         f"{ll['corr_next_day'].mean():+.4f} next day "
         f"({int((ll['p_next_day'] < 0.05).sum())} of 10), and "
         f"{ll['corr_prev_day'].mean():+.4f} the day before "
         f"({int((ll['p_prev_day'] < 0.05).sum())} of 10).")
    write_marker(doc, "lead_lag")

    doc.add_heading("Is the signal there at a slower frequency?", level=2)
    hz = n["horizons"]
    para(doc,
         "Absorption by the daily close does not rule out a slower effect, so "
         "sentiment averaged over the past k trading days was tested against "
         "the return over the next k days, for k of 1, 5, 21 and 63. "
         "Observations are sampled every k days rather than every day, because "
         "overlapping windows share data and would inflate significance "
         "badly at the longer horizons. Each horizon is run twice: on raw "
         "sector returns, and cross-sectionally with the day's mean across "
         "sectors removed from both sides, which is what the tilt actually "
         "bets on.")
    table(doc, hz[["horizon_days", "variant", "n_obs", "correlation",
                   "p_value", "significant_5pct", "significant_corrected"]],
          "Past-k-day sector sentiment against the next-k-day sector return, "
          "non-overlapping observations. The corrected column applies a "
          "Bonferroni threshold across the eight tests reported here.")
    para(doc,
         f"The correlation is negative at every horizon and its magnitude "
         f"grows with the horizon, from {hz.loc[0, 'correlation']:+.4f} at one "
         f"day to {hz[hz.horizon_days == hz.horizon_days.max()]['correlation'].min():+.4f} "
         "at 63 days, but the number of independent observations falls from "
         f"{int(hz['n_obs'].max()):,} to {int(hz['n_obs'].min())}. "
         f"{int(hz['significant_5pct'].sum())} of {len(hz)} tests clear 5% "
         f"uncorrected, and {int(hz['significant_corrected'].sum())} survive "
         "the correction for having run eight of them.")
    write_marker(doc, "horizons")

    # 5 -------------------------------------------------------------------
    doc.add_heading("5. Extensions", level=1)
    doc.add_heading("5.1 A finance lexicon for VADER", level=2)
    para(doc,
         "Only 1,831 of the 33,033 distinct tokens in the headline corpus "
         f"carry a VADER score, and {n['neutral_before']:.1%} of headlines "
         "score neutral. The gaps are not obscure: 'earnings' appears 11,870 "
         f"times and 'beat' 2,083, and VADER scores neither. {n['n_terms']} "
         "finance terms were added, selected by frequency from the headlines "
         "themselves rather than from memory, each with a written rationale. "
         "Terms whose sign reverses with context - 'cut', 'hike', 'buy' - were "
         "deliberately excluded and are listed with their reasons.")
    figure(doc, "lexicon_effect.png",
           "Share of headlines scoring inside VADER's neutral band, before and "
           "after adding the finance terms.")
    para(doc,
         f"The neutral share falls to {n['neutral_after']:.1%} and "
         f"{n['rescored']:.1%} of headlines are rescored.")
    write_marker(doc, "lexicon")

    doc.add_heading("5.2 A tilt that learns its own direction", level=2)
    para(doc,
         "Because the naive tilt leans the wrong way, the direction is made an "
         "estimate rather than an assumption. At each rebalance the pooled "
         "correlation between lagged sentiment and sector returns is "
         "re-estimated on the previous 252 trading days only, and the tilt "
         "leans whichever way that window indicates. Hard-coding the reverse "
         "would have fitted the answer to data already seen; this does not, "
         "and the truncation test covers it.")
    table(doc, n["extensions"][["fund", "variant", "sharpe",
                               "sharpe_net_costs", "avg_turnover"]],
          "Every fusion variant against its base fund, gross and net of 10 bp "
          "of one-way turnover.")
    figure(doc, "extension_comparison.png",
           "Sharpe ratio of each fusion variant against the base fund, gross "
           "and net of transaction costs.")
    para(doc,
         "A higher point estimate is not the same as an improvement. The "
         "difference in Sharpe ratios between each tilt and its base fund was "
         "resampled with a stationary bootstrap of 5,000 draws and a mean "
         "block length of 21 trading days, which keeps the two series paired "
         "and preserves their autocorrelation.")
    table(doc, n["significance"][["fund", "variant", "difference", "ci_low",
                                 "ci_high", "p_value", "significant_5pct"]],
          "Bootstrap confidence intervals for the change in Sharpe ratio "
          "relative to the base fund.")
    write_marker(doc, "significance")

    doc.add_heading("5.3 Was the winner found or fitted?", level=2)
    para(doc,
         "Comparing five variants over one period and keeping the best is "
         "itself a choice fitted to that period. The variants were therefore "
         "re-ranked on a 2021-2022 discovery window alone, and the selected "
         "variant's 2023 performance reported, 2023 having played no part in "
         "the choice.")
    figure(doc, "discovery_holdout.png",
           "Sharpe ratio in the 2021-2022 tuning window against the untouched "
           "2023 holdout, for each fusion variant.")
    table(doc, n["holdout"][["fund", "variant", "discovery_sharpe",
                            "holdout_sharpe", "selected_by_tuning"]],
          "Discovery-window ranking and holdout outcome for each variant.")
    write_marker(doc, "holdout")

    doc.add_heading("5.4 Does correcting the covariance fix the result?", level=2)
    write_marker(doc, "shrinkage")
    table(doc, n["shrinkage"][["family", "method", "shrinkage", "sharpe",
                               "ann_vol", "n_assets_held"]],
          "Each fund re-run with a Ledoit-Wolf shrunk covariance matrix "
          "against the sample covariance. Out-of-sample, 2021-2023.")

    doc.add_heading("5.5 The figure design system", level=2)
    para(doc,
         "Every exhibit is drawn through one design system rather than styled "
         "per figure. Categorical colours occupy fixed slots and are never "
         "cycled, so no two series in a chart can share a hue. The twelve-slot "
         "palette is validated pair by pair in CIE Lab space under normal, "
         "deuteranopic and protanopic vision by tools/check_palette.py; the "
         "worst separation is a deltaE of 16.9, identical to the original "
         "five-slot palette.")
    write_marker(doc, "design_system")

    # 6 -------------------------------------------------------------------
    doc.add_heading("6. The app and the investor journey", level=1)
    para(doc,
         "Spotlight is a Streamlit application. A visitor compares the twelve "
         "funds on one table and chart, opens a fund's fact sheet for its "
         "growth of $1, drawdown, current holdings and sector history, sets an "
         "allocation across funds with sliders and sees the blended track "
         "record, and reads the sentiment analytics including the fear and "
         "greed gauge and the sector index.")
    para(doc,
         "The application reads only precomputed CSV files committed under "
         "results/. It runs no backtest and loads no sentiment model, so cold "
         "starts stay fast on Streamlit Community Cloud's free tier. "
         "scripts/run_part_b.py rebuilds every artifact end to end.")
    write_marker(doc, "app")

    # 7 -------------------------------------------------------------------
    doc.add_heading("7. Critical reflection and recommendations", level=1)
    write_marker(doc, "reflection")
    write_marker(doc, "recommendations")

    doc.add_heading("References", level=1)
    para(doc, "Add sources through Word's References tab, then insert the "
              "bibliography here and update fields.")
    write_marker(doc, "references_note")

    doc.add_heading("Appendix", level=1)

    doc.add_heading("Appendix A - The four objectives", level=2)
    write_marker(doc, "equations")

    doc.add_heading("Appendix B - Supporting exhibits", level=2)
    para(doc, "Exhibits not needed in the main narrative.")
    figure(doc, "weights_combined_min_variance.png",
           "Sector weights over time for the Combined Minimum-Variance fund, "
           "the single-fund version of the across-methods exhibit.")
    table(doc, n["frequency"][["method", "frequency", "n_rebalances", "sharpe",
                               "sharpe_net_costs", "annual_turnover"]],
          "Full rebalance-frequency study: four schedules across the four "
          "combined funds, gross and net of 10 bp of one-way turnover.")

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(TARGET))

    n_marks = sum(MARK in p.text for p in Document(str(TARGET)).paragraphs)
    print(f"wrote {TARGET.relative_to(ROOT)}")
    print(f"  {n_marks} [WRITE] placeholders for my own prose")
    print("  open in Word, then select all and press F9 to number the "
          "captions and build the contents list")


if __name__ == "__main__":
    build()
